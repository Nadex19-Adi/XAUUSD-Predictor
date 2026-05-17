import docx

doc = docx.Document("skill report format.docx")

print("--- PARAGRAPHS ---")
for idx, para in enumerate(doc.paragraphs):
    if len(para.text.strip()) > 0:
        print(f"Para {idx}: {repr(para.text)}")

print("\n--- TABLES ---")
for idx, table in enumerate(doc.tables):
    print(f"Table {idx}: {len(table.rows)} rows, {len(table.columns)} columns")
    for r_idx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"  Row {r_idx}: {cells}")
